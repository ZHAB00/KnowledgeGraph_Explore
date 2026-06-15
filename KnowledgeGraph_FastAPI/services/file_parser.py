import io


class UnsupportedFormatError(Exception):
    def __init__(self, extension: str):
        self.extension = extension
        super().__init__(f"Unsupported file format: .{extension}")


class FileParser:
    SUPPORTED = {"txt", "md", "pdf", "docx"}

    def parse_from_bytes(self, data: bytes, filename: str) -> str:
        ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
        if ext not in self.SUPPORTED:
            raise UnsupportedFormatError(ext)

        if ext in ("txt", "md"):
            return data.decode("utf-8", errors="replace")
        elif ext == "pdf":
            return self._parse_pdf(data)
        elif ext == "docx":
            return self._parse_docx(data)
        return ""

    def _parse_pdf(self, data: bytes) -> str:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
                tables = page.extract_tables()
                for table in tables:
                    if table:
                        for row in table:
                            if row:
                                text_parts.append(" | ".join(
                                    str(c) if c else "" for c in row
                                ))
        return "\n\n".join(text_parts)

    def _parse_docx(self, data: bytes) -> str:
        from docx import Document
        doc = Document(io.BytesIO(data))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text for cell in row.cells]
                parts.append(" | ".join(cells))
        return "\n\n".join(parts)
